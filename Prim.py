import sys
import xlrd
import xlwt
import pandas as pd
import numpy as np

from docx import Document  #word相关
from docx.shared import Pt #word磅数

if __name__=='__main__':
    MAX = sys.maxsize
    # primgraph = [[MAX,  10, MAX, MAX, MAX,  11, MAX, MAX, MAX],
    #              [10,  MAX,  18, MAX, MAX, MAX,  16, MAX,  12],
    #              [MAX,  18, MAX,  22, MAX, MAX, MAX, MAX,   8],
    #              [MAX, MAX,  22, MAX,  20, MAX, MAX,  16,  21],
    #              [MAX, MAX, MAX,  20, MAX,  26,   7,  19, MAX],
    #              [11,  MAX, MAX, MAX,  26, MAX,  17, MAX, MAX],
    #              [MAX,  16, MAX, MAX,   7,  17, MAX,  19, MAX],
    #              [MAX, MAX, MAX,  16,  19, MAX,  19, MAX, MAX],
    #              [MAX,  12,   8,  21, MAX, MAX, MAX, MAX, MAX]]
    # chararray = ['A','B','C','D','E','F','G','H','I']


    out = [[-1 for i in range(30)] for j in range(30)]  # 这种创建方式，不会出现改其中的一个子list，其余全部都该的问题。

    # primgraph = [[MAX,  6, 1, 5, MAX, MAX],
    #              [6, MAX, 5, MAX,  3, MAX],
    #              [1, 5, MAX, 5, 6, 4],
    #              [5, MAX, 5, MAX, MAX, 2],
    #              [MAX, 3, 6, MAX, MAX, 6],
    #              [MAX, MAX, 4, 2, 6, MAX]]
    # chararray = ['A','B','C','D','E','F']

    data=[]
    table0 = xlrd.open_workbook(r'C:/Users/10947/Desktop/生活/cxl/MST/距离_2017.xlsx').sheet_by_index(0)  # 读取文件
    for i in range(1, 31):
        x1 = table0.col_values(i,1)
        for j in range(len(x1)):
            if (x1[j]=='N'):
                x1[j] = MAX
        # print(x1)
        data.append(x1)
    # print(data)
    chararray = table0.col_values(0,1)   # 最初给定的节点list
    # print(chararray)

    document = Document()

    charlist = []
    charlist.append(chararray[0])
    mid = []
    lowcost = []      # -1表示i已经在生成树集合中
    lowcost.append(-1)
    mid.append(0)
    n = len(chararray)
    for i in range(1,n):   # 初始化mid数组和lowcost数组
        # lowcost.append(primgraph[0][i])
        lowcost.append(data[0][i])
        mid.append(0)
    sum = 0
    for _ in range(1,n):   # 插入n-1个结点
        minid = 0
        min = MAX
        for j in range(1,n):  # 寻找每次插入生成树的权值最小的结点
            if(lowcost[j]!=-1 and lowcost[j]<min):
                minid = j
                min = lowcost[j]
        charlist.append(chararray[minid])  # minid：本轮选中的加入的点在最初给定的list中的序号，charlist：按顺序记录插入的节点名称
        print(chararray[mid[minid]]+'——'+chararray[minid]+'权值：'+str(lowcost[minid]))  #已在MST中的节点--新加入的节点+两个节点之间边的权重
        out[mid[minid]][minid]=lowcost[minid]
        out[minid][mid[minid]]=lowcost[minid]
        # document.add_paragraph(chararray[mid[minid]]+'——'+chararray[minid]+'权值：'+str(lowcost[minid]))
        sum+=min
        lowcost[minid] = -1

        for j in range(1,n):  # 这两个数组是混合数组，lowcost是将已经选中的所有节点对应的距离list，取这些list中最下的放到lowcost[]中，而mid[]则来记录lowcost[]中每个元素是从哪个节点对应的距离list中取出来的
            # if(lowcost[j]!=-1 and lowcost[j]>primgraph[minid][j]):
            #     lowcost[j] = primgraph[minid][j]
            if (lowcost[j] != -1 and lowcost[j] > data[minid][j]):
                lowcost[j] = data[minid][j]
                mid[j] = minid

    print("sum="+str(sum))
    print("插入结点顺序："+str(charlist))
    # document.add_paragraph("sum=" + str(sum))
    # document.add_paragraph("插入结点顺序："+str(charlist))
    # document.save('2017out.docx')
    print(out)

    # 测试代码
    # df = pd.DataFrame(out)
    # df.to_excel("textout2.xlsx", index=False)

    aux = [[0 for i in range(30)] for j in range(30)]  # 这种创建方式，不会出现改其中的一个子list，其余全部都该的问题。
    # 剪枝
    for i in range(n):
        for j in range(n):
            if(out[i][j] != -1 and aux[i][j] == 0):  # 邯郸i-长治j
                sum1 = 0
                sum2 = []  # 使用np计算均值和样本标准差
                num = 0
                aux[j][i] =1  # 邯郸-长治  与  长治-邯郸  是一致的，避免重复
                for m in range(n):
                    if(out[i][m]!=-1 and m!=j):  # 邯郸侧步长为1的边，另一节点为m
                        sum1 += out[i][m]
                        sum2.append(out[i][m])
                        num += 1
                        for p in range(n):
                            if (out[m][p] != -1 and p!=i):  # 邯郸侧步长为2的边，m-p，另一节点为p
                                sum1 += out[m][p]
                                sum2.append(out[m][p])
                                num += 1
                    if(out[j][m]!=-1 and m!=i): # 长治侧步长为1的边，另一节点为m
                        sum1 += out[j][m]
                        sum2.append(out[j][m])
                        num += 1
                        for p in range(n):
                            if (out[m][p] != -1 and p!=j):  # 长治侧步长为2的边，m-p，另一节点为p
                                sum1 += out[m][p]
                                sum2.append(out[m][p])
                                num += 1
                # avg = sum1/num
                avg = np.mean(sum2)
                std= np.std(sum2,ddof=1)
                # print("sum1="+str(sum1))
                # print(sum2)
                # print("num="+str(num))
                # print("avg="+str(avg))
                print(chararray[i]+'——'+chararray[j]+'的均值为' + str(avg))
                # print("std="+str(std))
                print(chararray[i]+'——'+chararray[j]+'的样本标准差为'+ str(std))
                document.add_paragraph(chararray[i] + '——' + chararray[j] + '的均值为' + str(avg))
                document.add_paragraph(chararray[i] + '——' + chararray[j] + '的样本标准差为' + str(std))
                if (out[i][j] > avg*2 or out[i][j] > std*3):
                    print(chararray[i]+'——'+chararray[j]+'的距离为'+str(out[i][j])+',为不连续边')
                    document.add_paragraph(chararray[i]+'——'+chararray[j]+'的距离为'+str(out[i][j])+',为不连续边')
                print('===============================================')
                document.add_paragraph('===============================================')
    document.save('2017out1.docx')

